from docx import Document
import os

os.chdir(r'C:\Users\28129\Desktop\共享1-数据收集与核对')

files = ['day1执行步骤.docx', '大致任务分工与时间规划.docx', '数据收集与核验问题.docx', '选题.doc']
for fname in files:
    print(f'===== {fname} =====')
    try:
        doc = Document(fname)
        for p in doc.paragraphs:
            if p.text.strip():
                print(p.text[:300])
        for t in doc.tables:
            for row in t.rows[:3]:
                cells = [c.text[:50] for c in row.cells]
                print(' | '.join(cells))
            print(f'  ... ({len(t.rows)} rows)')
    except Exception as e:
        print(f'Error: {e}')
    print()
